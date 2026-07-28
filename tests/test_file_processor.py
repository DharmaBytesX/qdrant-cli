import pytest

from qdrant_cli.file_processor import process_file, process_file_chunked


def test_process_file_unsupported_extension(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("dummy")
    with pytest.raises(ValueError, match="Unsupported file type"):
        process_file(str(f))


def test_process_file_not_found():
    with pytest.raises(FileNotFoundError):
        process_file("/nonexistent/file.docx")


@pytest.fixture
def sample_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello from docx test paragraph.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def multi_para_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Introduction to the technical concept.")
    doc.add_paragraph(
        "This is the second paragraph explaining details "
        "and deeper context about the system architecture and design choices."
    )
    doc.add_paragraph("A short third paragraph.")
    path = tmp_path / "multi.docx"
    doc.save(str(path))
    return str(path)


def test_process_file_docx(sample_docx):
    text = process_file(sample_docx)
    assert "Hello from docx test paragraph." in text


def test_process_file_chunked_docx(sample_docx):
    chunks = process_file_chunked(sample_docx, chunk_size=5)
    assert len(chunks) > 0
    assert all(isinstance(c, tuple) and len(c) == 2 for c in chunks)


def test_chunking_preserves_paragraph_boundaries(multi_para_docx):
    chunks = process_file_chunked(multi_para_docx, chunk_size=50)
    assert len(chunks) >= 1
    for text, _ in chunks:
        assert "\n\n" not in text


def test_chunking_small_paras_merged(multi_para_docx):
    chunks = process_file_chunked(multi_para_docx, chunk_size=100)
    para_1 = "Introduction to the technical concept."
    para_2 = (
        "This is the second paragraph explaining details "
        "and deeper context about the system architecture and design choices."
    )
    para_3 = "A short third paragraph."
    combined_text = " ".join(t for t, _ in chunks)
    assert para_1 in combined_text
    assert para_2 in combined_text
    assert para_3 in combined_text
